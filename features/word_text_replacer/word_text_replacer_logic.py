#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""features/word_text_replacer/word_text_replacer_logic.py — Word 批量文本替换逻辑层。

UI 与逻辑严格分离：本文件只负责文件扫描、docx 读取/替换/保存、结果汇总，
不依赖任何 Qt 类型。
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable


def _is_temp_docx(path: Path) -> bool:
    """Word 临时文件以 ~$ 开头，打开会报错，需要跳过。"""
    return path.name.startswith("~$")


def scan_docx_files(input_folder: Path, include_subfolders: bool = False) -> list[Path]:
    """扫描输入文件夹下所有 .docx 文件（排除 Word 临时文件）。

    Args:
        input_folder: 输入文件夹路径。
        include_subfolders: 是否递归包含子文件夹。

    Returns:
        按相对路径排序后的 .docx 文件路径列表。
    """
    if not input_folder.exists():
        return []

    pattern = "**/*.docx" if include_subfolders else "*.docx"
    files = [p for p in input_folder.glob(pattern)
             if p.is_file() and not _is_temp_docx(p)]
    files.sort(key=lambda p: p.relative_to(input_folder).as_posix())
    return files


def _replace_in_paragraph(paragraph, old: str, new: str) -> int:
    """在单个段落中替换文本，尽量保持原有格式。

    python-docx 的 run 可能把一段文字切得很碎（每个字一个 run），
    直接遍历 run.text 替换容易漏掉跨 run 的匹配。这里采用段落级整体文本替换：
    收集段落全部 text，执行替换后写回第一个 run，其余 run 清空。
    这样格式基本保留在原 run 上，且能处理跨 run 文本。

    Returns:
        该段落实际发生替换的次数。
    """
    if old == "":
        return 0

    full_text = paragraph.text
    if old not in full_text:
        return 0

    new_text = full_text.replace(old, new)
    count = full_text.count(old)

    runs = paragraph.runs
    if not runs:
        return 0

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""
    return count


def _replace_in_table(table, old: str, new: str) -> int:
    """递归替换表格所有单元格段落中的文本。"""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count += _replace_in_paragraph(paragraph, old, new)
            # 嵌套表格
            for nested in cell.tables:
                count += _replace_in_table(nested, old, new)
    return count


def replace_in_document(
    input_path: Path,
    output_path: Path,
    rules: list[tuple[str, str]],
) -> int:
    """打开 input_path 的 docx，按规则替换文本后保存到 output_path。

    Args:
        input_path: 源 docx 路径。
        output_path: 目标 docx 路径（父目录必须存在或会被创建）。
        rules: 替换规则列表，每项为 (old_text, new_text)。

    Returns:
        整个文档的总替换次数。
    """
    from docx import Document

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(input_path))

    total = 0
    for old, new in rules:
        if old == "":
            continue
        for paragraph in doc.paragraphs:
            total += _replace_in_paragraph(paragraph, old, new)
        for table in doc.tables:
            total += _replace_in_table(table, old, new)

    doc.save(str(output_path))
    return total


@dataclass
class ReplaceResult:
    """单次文件替换结果。"""
    input_path: Path
    output_path: Path
    replace_count: int
    ok: bool
    error: str | None = None


@dataclass
class BatchResult:
    """批量替换汇总结果。"""
    input_folder: Path
    output_folder: Path
    files: list[ReplaceResult]
    total_files: int
    success_count: int
    fail_count: int
    total_replacements: int
    log_lines: list[str]

    @property
    def ok(self) -> bool:
        return self.fail_count == 0


def process_folder(
    input_folder: Path,
    output_folder: Path,
    rules: list[tuple[str, str]],
    include_subfolders: bool = False,
    progress_callback: Callable[[str], None] | None = None,
) -> BatchResult:
    """批量替换入口。

    Args:
        input_folder: 输入文件夹。
        output_folder: 输出文件夹（必须不同于输入文件夹）。
        rules: 替换规则列表。
        include_subfolders: 是否递归处理子文件夹。
        progress_callback: 进度日志回调，用于向 UI 回传消息。

    Returns:
        BatchResult 汇总对象。
    """
    def _log(msg: str) -> None:
        if progress_callback is not None:
            progress_callback(msg)

    log_lines: list[str] = []

    # 校验
    if not input_folder or not output_folder:
        raise ValueError("输入文件夹和输出文件夹都必须选择")
    if input_folder == output_folder:
        raise ValueError("输入文件夹和输出文件夹不能相同（避免覆盖原文件）")
    if not input_folder.exists():
        raise FileNotFoundError(f"输入文件夹不存在：{input_folder}")

    files = scan_docx_files(input_folder, include_subfolders)
    _log(f"扫描到 {len(files)} 个 .docx 文件（{'含子文件夹' if include_subfolders else '仅当前文件夹'}）")

    if not files:
        return BatchResult(
            input_folder=input_folder,
            output_folder=output_folder,
            files=[],
            total_files=0,
            success_count=0,
            fail_count=0,
            total_replacements=0,
            log_lines=log_lines,
        )

    output_folder.mkdir(parents=True, exist_ok=True)

    results: list[ReplaceResult] = []
    success_count = 0
    fail_count = 0
    total_replacements = 0

    for src in files:
        rel = src.relative_to(input_folder)
        dst = output_folder / rel
        _log(f"处理：{rel.as_posix()}")
        try:
            count = replace_in_document(src, dst, rules)
            results.append(ReplaceResult(src, dst, count, ok=True))
            success_count += 1
            total_replacements += count
            _log(f"  → 替换 {count} 处，已保存到 {dst.relative_to(output_folder).as_posix()}")
        except Exception as e:  # noqa: BLE001
            results.append(ReplaceResult(src, dst, 0, ok=False, error=str(e)))
            fail_count += 1
            _log(f"  → 失败：{e}")

    return BatchResult(
        input_folder=input_folder,
        output_folder=output_folder,
        files=results,
        total_files=len(files),
        success_count=success_count,
        fail_count=fail_count,
        total_replacements=total_replacements,
        log_lines=log_lines,
    )
