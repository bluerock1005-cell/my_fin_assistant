#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""tests/test_word_text_replacer.py — word_text_replacer 逻辑层单元测试。

替换引擎已改为 COM / Word 自动化（见 word_text_replacer_logic.py），
因此涉及真实替换的用例需要本机安装 Microsoft Word；环境不具备时自动跳过。
纯文件系统扫描、输入/输出相同校验不依赖 Word，始终运行。
"""
from __future__ import annotations

import shutil
import tempfile
import unittest
from pathlib import Path

from features.word_text_replacer import word_text_replacer_logic as logic


def _word_available() -> bool:
    return logic._com_word_available()


# 仅在不具备 Word/COM 时跳过替换类用例
requires_word = unittest.skipUnless(_word_available(), "需要本机安装 Microsoft Word（COM）")


def _make_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


class TestScan(unittest.TestCase):
    def test_scan_docx_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "a.docx").write_text("")
            (root / "b.docx").write_text("")
            (root / "c.txt").write_text("")
            (root / "~$temp.docx").write_text("")
            sub = root / "sub"
            sub.mkdir()
            (sub / "d.docx").write_text("")

            found = logic.scan_docx_files(root, include_subfolders=False)
            self.assertEqual(len(found), 2)
            self.assertTrue(all(p.suffix == ".docx" for p in found))
            self.assertFalse(any(p.name.startswith("~$") for p in found))

            found_recursive = logic.scan_docx_files(root, include_subfolders=True)
            self.assertEqual(len(found_recursive), 3)


@requires_word
class TestReplace(unittest.TestCase):
    def test_replace_in_document(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "src.docx"
            dst = Path(tmp) / "dst.docx"
            _make_docx(src, ["Hello world", "world is great"])

            count = logic.replace_in_document(src, dst, [("world", "Python")])
            self.assertEqual(count, 2)

            from docx import Document
            doc = Document(str(dst))
            texts = [p.text for p in doc.paragraphs]
            self.assertEqual(texts, ["Hello Python", "Python is great"])

    def test_replace_in_table_header_footer(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "src.docx"
            dst = Path(tmp) / "dst.docx"
            from docx import Document
            d = Document()
            d.add_paragraph("Hello world in body")
            t = d.add_table(rows=1, cols=1)
            t.rows[0].cells[0].text = "world in table"
            d.sections[0].header.paragraphs[0].text = "world in header"
            d.sections[0].footer.paragraphs[0].text = "world in footer"
            d.save(str(src))

            count = logic.replace_in_document(src, dst, [("world", "PYTHON")])
            self.assertEqual(count, 4)

            out = Document(str(dst))
            self.assertEqual(out.paragraphs[0].text, "Hello PYTHON in body")
            self.assertEqual(out.tables[0].rows[0].cells[0].text, "PYTHON in table")
            self.assertEqual(out.sections[0].header.paragraphs[0].text, "PYTHON in header")
            self.assertEqual(out.sections[0].footer.paragraphs[0].text, "PYTHON in footer")

    def test_process_folder(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_folder = root / "input"
            output_folder = root / "output"
            input_folder.mkdir()
            output_folder.mkdir()

            _make_docx(input_folder / "a.docx", ["foo bar"])
            _make_docx(input_folder / "b.docx", ["bar baz"])

            logs: list[str] = []
            result = logic.process_folder(
                input_folder,
                output_folder,
                [("bar", "BAR")],
                include_subfolders=False,
                progress_callback=logs.append,
            )

            self.assertEqual(result.total_files, 2)
            self.assertEqual(result.success_count, 2)
            self.assertEqual(result.fail_count, 0)
            self.assertEqual(result.total_replacements, 2)
            self.assertGreater(len(logs), 0)

            from docx import Document
            self.assertTrue((output_folder / "a.docx").exists())
            doc = Document(str(output_folder / "a.docx"))
            self.assertEqual(doc.paragraphs[0].text, "foo BAR")


class TestValidation(unittest.TestCase):
    def test_input_equals_output_raises(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            folder = Path(tmp)
            with self.assertRaises(ValueError) as ctx:
                logic.process_folder(folder, folder, [("a", "b")])
            self.assertIn("不能相同", str(ctx.exception))


if __name__ == "__main__":
    unittest.main(verbosity=2)
